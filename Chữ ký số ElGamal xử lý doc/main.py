import tkinter as tk
from tkinter import filedialog
from tkinter import scrolledtext
from tkinter import messagebox
from tkinter import ttk
from docx import Document

from signElGamal import sign, verify

import keyElGamal
from keyElGamal import generate_key, is_prime

# Khởi tạo cửa sổ giao diện
window = tk.Tk()
window.title("Nhóm 14 - ATBMTT")

notebook = ttk.Notebook(window, style='TNotebook')
notebook.pack(fill='both', expand=True)

sign_tab = ttk.Frame(notebook)
notebook.add(sign_tab, text="Ký")

file_imported = False
path_file_to_sign = ""

def sign_button_click():
    if file_imported:
        print(f"Ký văn bản Word: {path_file_to_sign}")
        sign_code = sign(path_file_to_sign, is_file=True)
    else:
        text = text_entry.get("1.0", "end-1c")
        if text == "":
            messagebox.showerror("Lỗi", "Vui lòng nhập vào ô văn bản!")
            return
        print("Ký văn bản đã nhập:", text)
        sign_code = sign(text, is_file=False)

    if sign_code is None:
        messagebox.showerror("Lỗi", "Không thể tạo chữ ký! Vui lòng kiểm tra khóa hoặc dữ liệu đầu vào.")
        return

    sign_entry.config(state=tk.NORMAL)
    sign_entry.delete("1.0", "end")
    sign_entry.insert("1.0", sign_code)
    sign_entry.config(state=tk.DISABLED)

def import_button_click():
    global path_file_to_sign
    global file_imported
    test_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx"), ("All Files", "*.*")])

    if test_path:
        if file_imported:
            cancel_button_click()
        path_file_to_sign = test_path
        document = Document(path_file_to_sign)
        content = "\n".join([paragraph.text for paragraph in document.paragraphs])
        text_entry.delete("1.0", "end")
        text_entry.insert("1.0", content)
        text_entry.config(state=tk.DISABLED)
        file_imported = True
        messagebox.showinfo("Thông báo", "Import file thành công!")
        info_label.config(text="Đã import tệp tin Word: " + path_file_to_sign)
        info_label.pack(pady=(10, 0))
        cancel_button.pack(pady=(10, 0), side=tk.BOTTOM)

def forward_button_click():
    global path_file_to_verify, right_file_imported
    if right_file_imported:
        right_cancel_button_click()

    if file_imported:
        path_file_to_verify = path_file_to_sign
        if path_file_to_verify:
            document = Document(path_file_to_verify)
            content = "\n".join([paragraph.text for paragraph in document.paragraphs])
            right_text_entry.delete("1.0", "end")
            right_text_entry.insert("1.0", content)
            right_text_entry.config(state=tk.DISABLED)
            right_file_imported = True
            right_info_label.config(text="Đã import tệp tin Word: " + path_file_to_verify)
            right_info_label.pack(pady=(10, 0))
            right_cancel_button.pack(pady=(10, 0), side=tk.BOTTOM)
    else:
        temp = text_entry.get("1.0", "end-1c")
        if temp:
            right_text_entry.delete("1.0", "end")
            right_text_entry.insert("1.0", temp)

    temp = sign_entry.get("1.0", "end-1c")
    if temp:
        right_sign_entry.delete("1.0", "end")
        right_sign_entry.insert("1.0", temp)

def save_button_click():
    text = sign_entry.get("1.0", "end-1c")
    if text:
        save_file_path = filedialog.asksaveasfilename(defaultextension=".txt")
        if save_file_path:
            with open(save_file_path, "w") as file:
                file.write(text)
            messagebox.showinfo("Thông báo", "Lưu chữ ký thành công!")
            print("Lưu Chữ ký thành công! Đường dẫn:", save_file_path)

def cancel_button_click():
    text_entry.config(state=tk.NORMAL)
    text_entry.delete("1.0", "end")
    info_label.config(text="")
    info_label.pack(pady=0)
    cancel_button.pack_forget()
    global file_imported
    file_imported = False

left_frame = tk.LabelFrame(sign_tab, text="Phát sinh chữ ký")
left_frame.pack(side=tk.LEFT, padx=10, pady=10, anchor="nw")

# Ô VĂN BẢN KÝ
text_entry_frame = tk.Frame(left_frame)
text_entry_frame.pack(side=tk.TOP)

text_entry_label = tk.Label(text_entry_frame, text="Văn bản ký:", font=("Arial", 12))
text_entry_label.pack(padx=20, pady=(10, 0), anchor="nw")

scrollbar1 = tk.Scrollbar(text_entry_frame)
text_entry = scrolledtext.ScrolledText(text_entry_frame, height=15, width=50, yscrollcommand=scrollbar1.set, font=("Arial", 10))
text_entry.pack(padx=20, side=tk.LEFT, fill=tk.BOTH)
scrollbar1.config(command=text_entry.yview)

text_entry_buttons_frame = tk.Frame(text_entry_frame)
text_entry_buttons_frame.pack(padx=(0, 20), side=tk.RIGHT)
import_button = tk.Button(text_entry_buttons_frame, text="Nhập file Word", command=import_button_click, bg="#19A7CE")
import_button.pack(side=tk.TOP)
cancel_button = tk.Button(text_entry_buttons_frame, text="Hủy", command=cancel_button_click)
# GIỮA 2 Ô
info_label = tk.Label(left_frame, text="", wraplength=500, font=("Arial", 10))
info_label.pack(padx=20)

sign_button = tk.Button(left_frame, text="Ký văn bản", command=sign_button_click, bg="#19A7CE")
sign_button.pack(pady=(10, 0))
# Ô CHỮ KÝ
sign_entry_frame = tk.Frame(left_frame)
sign_entry_frame.pack(side=tk.BOTTOM)

sign_entry_label = tk.Label(sign_entry_frame, text="Chữ ký:", font=("Arial", 12))
sign_entry_label.pack(pady=(10, 0), anchor="nw")

scrollbar2 = tk.Scrollbar(sign_entry_frame)
sign_entry = scrolledtext.ScrolledText(sign_entry_frame, height=5, width=50, yscrollcommand=scrollbar2.set, font=("Arial", 10))
sign_entry.pack(padx=(0, 20), pady=(0, 20), side=tk.LEFT, fill=tk.BOTH)
sign_entry.config(state=tk.DISABLED)
scrollbar2.config(command=sign_entry.yview)

sign_buttons_frame = tk.Frame(sign_entry_frame)
sign_buttons_frame.pack(padx=20, side=tk.RIGHT)

forward_button = tk.Button(sign_buttons_frame, text="Chuyển", command=forward_button_click, bg="#19A7CE")
forward_button.pack(pady=(0, 10), side=tk.TOP)

save_button = tk.Button(sign_buttons_frame, text="Lưu", command=save_button_click, bg="#19A7CE")
save_button.pack(side=tk.BOTTOM)

right_frame = tk.LabelFrame(sign_tab, text="Xác nhận chữ ký")
right_frame.pack(side=tk.RIGHT, padx=(0, 10), pady=10)
right_file_imported = False
path_file_to_verify = ""

def verify_sign_button_click():
    sign_temp = right_sign_entry.get("1.0", "end-1c")
    if sign_temp is None or sign_temp == "":
        messagebox.showerror("Lỗi", "Vui lòng nhập chữ ký!")
        return

    if right_file_imported:
        print(f"Kiểm tra văn bản Word: {path_file_to_verify}")
        correct_sign = verify(path_file_to_verify, sign_temp, right_file_imported)
    else:
        text = right_text_entry.get("1.0", "end-1c")
        if text == "":
            messagebox.showerror("Lỗi", "Vui lòng nhập vào ô văn bản!")
            return
        print("Kiểm tra văn bản đã nhập:", text)
        correct_sign = verify(text, sign_temp, right_file_imported)

    if correct_sign is None:
        messagebox.showerror("Lỗi", "Chữ ký không hợp lệ hoặc dữ liệu đầu vào không đúng!")
        return

    sign_verifying_state_text.config(state=tk.NORMAL)
    if correct_sign:
        sign_verifying_state_text.delete("1.0", "end")
        sign_verifying_state_text.insert("1.0", "Chữ ký chính xác!")
        print("Chữ ký chính xác!")
    else:
        sign_verifying_state_text.delete("1.0", "end")
        sign_verifying_state_text.insert("1.0", "Chữ ký không chính xác!")
        print("Chữ ký không chính xác!")
    sign_verifying_state_text.config(state=tk.DISABLED)

def right_import_button_click():
    global path_file_to_verify
    global right_file_imported
    test_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx"), ("All Files", "*.*")])

    if test_path:
        if right_file_imported:
            right_cancel_button_click()
        path_file_to_verify = test_path
        document = Document(path_file_to_verify)
        content = "\n".join([paragraph.text for paragraph in document.paragraphs])
        right_text_entry.delete("1.0", "end")
        right_text_entry.insert("1.0", content)
        right_text_entry.config(state=tk.DISABLED)
        right_file_imported = True
        messagebox.showinfo("Thông báo", "Import file thành công!")
        right_info_label.config(text="Đã import tệp tin Word: " + path_file_to_verify)
        right_info_label.pack(pady=(10, 0))
        right_cancel_button.pack(pady=(10, 0), side=tk.BOTTOM)

def right_sign_import_file_click():
    path_file_sign = filedialog.askopenfilename(filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")])
    if path_file_sign:
        with open(path_file_sign, "r") as f:
            right_sign_entry.delete("1.0", "end")
            right_sign_entry.insert("1.0", f.read())

def right_cancel_button_click():
    right_text_entry.config(state=tk.NORMAL)
    right_text_entry.delete("1.0", "end")
    right_info_label.config(text="")
    right_info_label.pack(pady=0)
    right_cancel_button.pack_forget()
    global right_file_imported
    right_file_imported = False

# Ô VĂN BẢN
text_entry_frame = tk.Frame(right_frame)
text_entry_frame.pack()

text_entry_label = tk.Label(text_entry_frame, text="Văn bản xác nhận:", font=("Arial", 12))
text_entry_label.pack(padx=20, pady=(10, 0), anchor="nw")

scrollbar1 = tk.Scrollbar(text_entry_frame)
right_text_entry = scrolledtext.ScrolledText(text_entry_frame, height=15, width=50, yscrollcommand=scrollbar1.set, font=("Arial", 10))
right_text_entry.pack(padx=20, side=tk.LEFT, fill=tk.BOTH)
scrollbar1.config(command=right_text_entry.yview)

text_entry_buttons_frame = tk.Frame(text_entry_frame)
text_entry_buttons_frame.pack(padx=(0, 20), side=tk.RIGHT)
import_button = tk.Button(text_entry_buttons_frame, text="Nhập file Word", command=right_import_button_click, bg="#19A7CE")
import_button.pack(side=tk.TOP)
right_cancel_button = tk.Button(text_entry_buttons_frame, text="Hủy", command=right_cancel_button_click)
# GIỮA 2 Ô
right_info_label = tk.Label(right_frame, text="", wraplength=500, font=("Arial", 10))
right_info_label.pack(padx=20)

verify_sign_button = tk.Button(right_frame, text="Kiểm tra chữ ký", command=verify_sign_button_click, bg="#19A7CE")
verify_sign_button.pack(pady=(10, 0))
# Ô CHỮ KÝ
right_sign_entry_frame = tk.Frame(right_frame)
right_sign_entry_frame.pack()

right_sign_entry_label = tk.Label(right_sign_entry_frame, text="Chữ ký:", font=("Arial", 12))
right_sign_entry_label.pack(padx=20, pady=(10, 0), anchor="nw")

scrollbar2 = tk.Scrollbar(right_sign_entry_frame)
right_sign_entry = scrolledtext.ScrolledText(right_sign_entry_frame, height=5, width=50, yscrollcommand=scrollbar2.set, font=("Arial", 10))
right_sign_entry.pack(padx=20, pady=(0, 20), side=tk.LEFT, fill=tk.BOTH)
scrollbar2.config(command=right_sign_entry.yview)

verify_sign_buttons_frame = tk.Frame(right_sign_entry_frame)
verify_sign_buttons_frame.pack(side=tk.RIGHT)
import_sign_file_button = tk.Button(verify_sign_buttons_frame, text="Nhập file chữ ký", command=right_sign_import_file_click, bg="#19A7CE")
import_sign_file_button.pack(padx=(0, 20))
# Ô THÔNG BÁO
sign_verifying_state_frame = tk.Frame(right_frame)
sign_verifying_state_frame.pack(anchor="sw")

sign_verifying_state_label = tk.Label(sign_verifying_state_frame, text="Thông báo:", font=("Arial", 12))
sign_verifying_state_label.pack(padx=20, pady=(10, 0), anchor="nw")

sign_verifying_state_text = tk.Text(sign_verifying_state_frame, width=50, height=4, font=("Arial", 10))
sign_verifying_state_text.pack(padx=(20, 0), pady=(0, 20), side=tk.LEFT)
sign_verifying_state_text.config(state=tk.DISABLED)

# TAB TẠO KHÓA
def only_numbers(char):
    return char.isdigit()

validation = window.register(only_numbers)

def getKeyVal():
    entry_p.config(state=tk.NORMAL)
    entry_p.delete("0", tk.END)
    entry_p.insert(tk.END, keyElGamal.PUBLIC_KEY[0])
    entry_p.config(state=tk.DISABLED)

    entry_g.config(state=tk.NORMAL)
    entry_g.delete("0", tk.END)
    entry_g.insert(tk.END, keyElGamal.PUBLIC_KEY[1])
    entry_g.config(state=tk.DISABLED)

    entry_y.config(state=tk.NORMAL)
    entry_y.delete("0", tk.END)
    entry_y.insert(tk.END, keyElGamal.PUBLIC_KEY[2])
    entry_y.config(state=tk.DISABLED)

    entry_x.config(state=tk.NORMAL)
    entry_x.delete("0", tk.END)
    entry_x.insert(tk.END, keyElGamal.PRIVATE_KEY[2])
    entry_x.config(state=tk.DISABLED)

def generate_key_button_click():
    p_new = int(entry_p.get())
    if p_new > 1000:
        messagebox.showerror("Lỗi", "P quá lớn!")
        return
    if not is_prime(p_new):
        messagebox.showerror("Lỗi", "P phải là số nguyên tố!")
        return

    keyElGamal.PUBLIC_KEY, keyElGamal.PRIVATE_KEY = generate_key(p_new)
    messagebox.showinfo("Thông báo", "Tạo khóa mới thành công!")
    getKeyVal()

key_tab = ttk.Frame(notebook)
notebook.add(key_tab, text="Khóa")

p_frame = ttk.LabelFrame(key_tab, text="Chọn số nguyên tố")
p_frame.pack(side=tk.TOP, padx=10, pady=30)

label_p = ttk.Label(p_frame, text="P:", font=("Arial", 13))
label_p.grid(row=0, column=0, padx=(50, 10), pady=(30, 20))
entry_p = ttk.Entry(p_frame, width=10, validate="key", validatecommand=(validation, '%S'), font=("Arial", 13))
entry_p.grid(row=0, column=1, padx=(10, 50), pady=(30, 20))
entry_p.insert(tk.END, str(keyElGamal.PUBLIC_KEY[0]))

key_frame = ttk.Frame(key_tab)
key_frame.pack(side=tk.BOTTOM, pady=(20, 50))

pubKey_frame = ttk.LabelFrame(key_tab, text="Khóa công khai")
pubKey_frame.pack(side=tk.LEFT, anchor="nw", padx=(60, 0), pady=(50, 20))

label_g = ttk.Label(pubKey_frame, text="G:", font=("Arial", 13))
label_g.grid(row=0, column=0, padx=(50, 10), pady=(30, 20))
entry_g = tk.Entry(pubKey_frame, width=10, font=("Arial", 13), disabledbackground="white")
entry_g.grid(row=0, column=1, padx=(10, 30), pady=(30, 20))

label_y = ttk.Label(pubKey_frame, text="Y:", font=("Arial", 13))
label_y.grid(row=1, column=0, padx=(50, 10), pady=(20, 40))
entry_y = tk.Entry(pubKey_frame, width=10, font=("Arial", 13), disabledbackground="white")
entry_y.grid(row=1, column=1, padx=(10, 30), pady=(20, 40))

label_p_pub = ttk.Label(pubKey_frame, text="P:", font=("Arial", 13))
label_p_pub.grid(row=2, column=0, padx=(50, 10), pady=(20, 40))
entry_p_pub = tk.Entry(pubKey_frame, width=10, font=("Arial", 13), disabledbackground="white")
entry_p_pub.grid(row=2, column=1, padx=(10, 30), pady=(20, 40))

priKey_frame = ttk.LabelFrame(key_tab, text="Khóa bí mật")
priKey_frame.pack(side=tk.RIGHT, anchor="ne", padx=(0, 60), pady=(50, 20))

label_x = ttk.Label(priKey_frame, text="X:", font=("Arial", 13))
label_x.grid(row=0, column=0, padx=(50, 10), pady=(30, 40))
entry_x = tk.Entry(priKey_frame, width=10, font=("Arial", 13), disabledbackground="white")
entry_x.grid(row=0, column=1, padx=(10, 50), pady=(30, 40))

getKeyVal()
generate_button = tk.Button(key_tab, text="Tạo", font=("Arial", 14), command=generate_key_button_click, bg="#19A7CE")
generate_button.pack(side=tk.BOTTOM, pady=(30, 0), ipadx=20, ipady=10)

window.mainloop()